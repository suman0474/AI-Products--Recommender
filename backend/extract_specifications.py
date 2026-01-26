#!/usr/bin/env python3
"""
Extract technical specifications from DOCX files in chroma_data/Standards
and organize by product type into a structured folder hierarchy.
"""

import os
import json
import re
from pathlib import Path
from docx import Document
from typing import Dict, List, Any
import logging

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Define product type mapping from filename
PRODUCT_TYPE_MAPPING = {
    'pressure': 'Pressure Measurement',
    'temperature': 'Temperature Measurement',
    'flow': 'Flow Measurement',
    'level': 'Level Measurement',
    'control_systems': 'Control Systems',
    'valves_actuators': 'Valves & Actuators',
    'calibration_maintenance': 'Calibration & Maintenance',
    'comm_signal': 'Communication & Signals',
    'condition_monitoring': 'Condition Monitoring',
    'analytical': 'Analytical Instrumentation',
    'safety': 'Safety & Protection',
    'accessories_calibration': 'Accessories & Calibration',
}

def extract_text_from_docx(filepath: str) -> str:
    """Extract all text from a DOCX file."""
    try:
        doc = Document(filepath)
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        return '\n'.join(text_content)
    except Exception as e:
        logger.error(f"Error extracting text from {filepath}: {e}")
        return ""

def extract_specifications(text: str) -> Dict[str, Any]:
    """
    Extract technical specifications from document text.
    Looks for patterns like:
    - Standard codes (IEC, ISO, API, ANSI, ISA)
    - Certifications (SIL, ATEX, CE, UL)
    - Measurement ranges
    - Accuracy
    - Response time
    """
    specs = {
        'standards': [],
        'certifications': [],
        'measurement_ranges': [],
        'accuracy_specs': [],
        'response_time': [],
        'environmental_specs': [],
        'communication_protocols': [],
        'safety_requirements': [],
        'calibration_info': [],
        'raw_text_preview': text[:1000] if len(text) > 1000 else text
    }

    # Extract standard codes
    standard_patterns = [
        (r'\b(IEC\s*\d+(?:[.-]\d+)*(?:\s*Part\s*\d+)?)\b', 'IEC'),
        (r'\b(ISO\s*\d+(?:[.-]\d+)*)\b', 'ISO'),
        (r'\b(API\s*\d+(?:[.-]\d+)*)\b', 'API'),
        (r'\b(ANSI[/\s]+\w+\s*\d+(?:[.-]\d+)*)\b', 'ANSI'),
        (r'\b(ISA[/\s]*\d+(?:[.-]\d+)*)\b', 'ISA'),
        (r'\b(EN\s*\d+(?:[.-]\d+)*)\b', 'EN'),
        (r'\b(NFPA\s*\d+)\b', 'NFPA'),
        (r'\b(ASME\s*[A-Z]*\s*\d+(?:[.-]\d+)*)\b', 'ASME'),
    ]

    for pattern, standard_type in standard_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            specs['standards'].append({
                'code': match.strip(),
                'type': standard_type
            })

    # Extract certifications
    cert_patterns = [
        (r'\bSIL\s*[1-4]\b', 'SIL'),
        (r'\bATEX\s*(?:Zone\s*)?[0-2]?(?:/[0-2]+)?\b', 'ATEX'),
        (r'\bIECEx\b', 'IECEx'),
        (r'\bCE\s*(?:Mark(?:ed)?)?\b', 'CE Mark'),
        (r'\b(UL|cUL|cULus)\b', 'UL'),
        (r'\bCSA\b', 'CSA'),
        (r'\b(IP\d{2})\b', 'IP Rating'),
        (r'\bNEMA\s*[0-9]+[A-Z]*\b', 'NEMA'),
        (r'\bFM\s*(?:Approved)?\b', 'FM'),
        (r'\bClass\s*[I]+,?\s*Div(?:ision)?\s*[1-2]\b', 'Class I Div'),
    ]

    for pattern, cert_type in cert_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            specs['certifications'].append({
                'name': match.strip().upper(),
                'type': cert_type
            })

    # Extract measurement ranges
    range_patterns = [
        (r'(?:pressure|measurement)\s*(?:range|:\s*)?(-?\d+\.?\d*)\s*(?:to|[-–])\s*(-?\d+\.?\d*)\s*(bar|psi|MPa|kPa|Pa)', 'pressure'),
        (r'(?:temperature|temp)\s*(?:range|:\s*)?(-?\d+\.?\d*)\s*(?:to|[-–])\s*(-?\d+\.?\d*)\s*(?:°C|°F|K|degrees)', 'temperature'),
        (r'(?:flow)\s*(?:range|:\s*)?(\d+\.?\d*)\s*(?:to|[-–])\s*(\d+\.?\d*)\s*(m³/h|l/min|GPM|SCFM)', 'flow'),
    ]

    for pattern, range_type in range_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            specs['measurement_ranges'].append({
                'type': range_type,
                'min': match[0],
                'max': match[1],
                'unit': match[2] if len(match) > 2 else 'unknown'
            })

    # Extract accuracy
    accuracy_patterns = [
        r'(?:accuracy|precision)[:\s]+([±]?\s*\d+\.?\d*\s*%)',
        r'([±]?\s*\d+\.?\d*\s*%)\s*(?:of|full|scale)',
    ]

    for pattern in accuracy_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            specs['accuracy_specs'].append(match.strip())

    # Extract response time
    response_patterns = [
        r'(?:response\s*time)[:\s]+(\d+\.?\d*\s*(?:ms|seconds?|s))',
        r'(?:settling\s*time)[:\s]+(\d+\.?\d*\s*(?:ms|seconds?|s))',
    ]

    for pattern in response_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            specs['response_time'].append(match.strip())

    # Extract communication protocols
    protocols = re.findall(
        r'\b(4-20\s*mA|HART|Profibus|Foundation\s*Fieldbus|Modbus|EtherNet/IP|PROFINET|IO-Link|Ethernet|analog|digital)\b',
        text,
        re.IGNORECASE
    )
    specs['communication_protocols'] = list(set([p.strip().upper() for p in protocols]))

    # Extract environmental specs
    env_patterns = [
        r'(?:storage|operating)\s*(?:temperature|temp)[:\s]+([-\d\s°CFK/–to]+)',
        r'(?:humidity)[:\s]+(\d+\s*[-–]\s*\d+\s*%)',
        r'(?:IP\s*rating)[:\s]+(IP\d{2})',
    ]

    for pattern in env_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        specs['environmental_specs'].extend([m.strip() for m in matches])

    # Extract safety requirements
    if any(keyword in text.lower() for keyword in ['sil', 'safety', 'functional safety', 'atex', 'hazardous']):
        specs['safety_requirements'].append('Safety-related requirements present')

    # Extract calibration info
    if any(keyword in text.lower() for keyword in ['calibration', 'traceability', 'nist', 'uncertainty']):
        specs['calibration_info'].append('Calibration requirements present')

    return specs

def analyze_chroma_data(chroma_path: str) -> Dict[str, Any]:
    """Analyze all documents in chroma_data directory."""
    standards_dir = os.path.join(chroma_path, 'Standards')

    if not os.path.exists(standards_dir):
        logger.error(f"Standards directory not found: {standards_dir}")
        return {}

    analysis_results = {}
    docx_files = [f for f in os.listdir(standards_dir) if f.endswith('.docx')]

    logger.info(f"Found {len(docx_files)} DOCX files in {standards_dir}")

    for filename in docx_files:
        filepath = os.path.join(standards_dir, filename)
        logger.info(f"Processing: {filename}")

        # Determine product type
        product_type = None
        for key, value in PRODUCT_TYPE_MAPPING.items():
            if key in filename.lower():
                product_type = value
                break

        if not product_type:
            product_type = 'Other'

        # Extract text
        text = extract_text_from_docx(filepath)

        # Extract specifications
        specs = extract_specifications(text)

        if product_type not in analysis_results:
            analysis_results[product_type] = []

        analysis_results[product_type].append({
            'source_file': filename,
            'specifications': specs,
            'text_length': len(text),
            'paragraphs': len(text.split('\n'))
        })

    return analysis_results

def create_folder_structure(output_base: str, analysis: Dict[str, Any]) -> None:
    """Create folder structure and save organized specifications."""

    # Create base output directory
    os.makedirs(output_base, exist_ok=True)
    logger.info(f"Created output directory: {output_base}")

    # Create subdirectories for each product type
    for product_type, documents in analysis.items():
        product_dir = os.path.join(output_base, product_type)
        os.makedirs(product_dir, exist_ok=True)

        logger.info(f"Created directory for: {product_type}")

        # Save specifications summary
        summary_file = os.path.join(product_dir, '00_SUMMARY.json')
        summary_data = {
            'product_type': product_type,
            'total_documents': len(documents),
            'documents': documents,
            'aggregated_specs': aggregate_specifications(documents)
        }

        with open(summary_file, 'w', encoding='utf-8') as f:
            json.dump(summary_data, f, indent=2, ensure_ascii=False)

        logger.info(f"Saved summary: {summary_file}")

        # Save individual document specs
        for idx, doc_info in enumerate(documents, 1):
            spec_file = os.path.join(
                product_dir,
                f"{idx:02d}_{doc_info['source_file'].replace('.docx', '.json')}"
            )
            with open(spec_file, 'w', encoding='utf-8') as f:
                json.dump(doc_info, f, indent=2, ensure_ascii=False)

        # Create README
        readme_file = os.path.join(product_dir, 'README.txt')
        with open(readme_file, 'w', encoding='utf-8') as f:
            f.write(f"Product Type: {product_type}\n")
            f.write(f"Total Documents: {len(documents)}\n")
            f.write(f"===========================================\n\n")
            for doc in documents:
                f.write(f"Source: {doc['source_file']}\n")
                f.write(f"Text Length: {doc['text_length']} characters\n")
                f.write(f"Paragraphs: {doc['paragraphs']}\n\n")

def aggregate_specifications(documents: List[Dict]) -> Dict[str, Any]:
    """Aggregate specifications across all documents of a product type."""
    aggregated = {
        'total_standards': [],
        'total_certifications': [],
        'total_protocols': [],
        'document_count': len(documents)
    }

    for doc in documents:
        specs = doc.get('specifications', {})
        aggregated['total_standards'].extend(specs.get('standards', []))
        aggregated['total_certifications'].extend(specs.get('certifications', []))
        aggregated['total_protocols'].extend(specs.get('communication_protocols', []))

    # Deduplicate
    aggregated['total_standards'] = list({json.dumps(s, sort_keys=True) for s in aggregated['total_standards']})
    aggregated['total_standards'] = [json.loads(s) for s in aggregated['total_standards']]

    aggregated['total_certifications'] = list(set(json.dumps(c, sort_keys=True) for c in aggregated['total_certifications']))
    aggregated['total_certifications'] = [json.loads(c) for c in aggregated['total_certifications']]

    aggregated['total_protocols'] = list(set(aggregated['total_protocols']))

    return aggregated

if __name__ == '__main__':
    chroma_data_path = r'D:\AI PR\AIPR\backend\chroma_data'
    output_path = r'D:\AI PR\AIPR\backend\PRODUCT_SPECIFICATIONS'

    logger.info("Starting specification extraction and organization...")
    logger.info(f"Source: {chroma_data_path}")
    logger.info(f"Output: {output_path}")

    # Analyze chroma data
    analysis = analyze_chroma_data(chroma_data_path)

    if analysis:
        logger.info(f"Found {len(analysis)} product types")

        # Create folder structure
        create_folder_structure(output_path, analysis)

        logger.info("=" * 60)
        logger.info("EXTRACTION COMPLETE")
        logger.info("=" * 60)

        # Print summary
        for product_type in sorted(analysis.keys()):
            docs = analysis[product_type]
            logger.info(f"\n{product_type}:")
            logger.info(f"  - Documents: {len(docs)}")

            # Aggregate specs
            all_standards = []
            all_certs = []
            for doc in docs:
                all_standards.extend(doc['specifications'].get('standards', []))
                all_certs.extend(doc['specifications'].get('certifications', []))

            unique_standards = list(set(json.dumps(s, sort_keys=True) for s in all_standards))
            unique_certs = list(set(json.dumps(c, sort_keys=True) for c in all_certs))

            logger.info(f"  - Unique Standards Found: {len(unique_standards)}")
            logger.info(f"  - Unique Certifications: {len(unique_certs)}")
    else:
        logger.error("No documents found to analyze")
