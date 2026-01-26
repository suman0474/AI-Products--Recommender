"""
File Extraction Utilities
Extracts text from various file formats: PDF, DOCX, TXT, Images (OCR via Gemini Vision)
"""

import logging
import io
import os
import fitz  # PyMuPDF for PDF
from docx import Document  # python-docx for DOCX
from PIL import Image
import google.generativeai as genai

# Initialize logger
logger = logging.getLogger(__name__)

# Initialize Gemini Client for OCR
try:
    api_key = os.getenv("GOOGLE_API_KEY")
    if api_key:
        genai.configure(api_key=api_key)
        logger.info("Gemini Vision initialized for OCR fallback")
    else:
        logger.warning("GOOGLE_API_KEY not found. Image OCR will not work.")
except Exception as e:
    logger.error(f"Failed to initialize Gemini for OCR: {e}")

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file"""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page_num, page in enumerate(doc):
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
        doc.close()
        extracted_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(extracted_text)} characters from PDF")
        return extracted_text
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {e}")
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    text_parts.append(row_text)
        extracted_text = "\n".join(text_parts)
        logger.info(f"Extracted {len(extracted_text)} characters from DOCX")
        return extracted_text
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        return ""

def extract_text_from_image(file_bytes: bytes) -> str:
    """
    Extract text from image using Gemini Vision (Cloud OCR)
    Replaces heavy local EasyOCR dependency.
    """
    try:
        image = Image.open(io.BytesIO(file_bytes))
        
        # Use Gemini 1.5 Flash for fast OCR
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        response = model.generate_content([
            "Transcribe all text from this image exactly as it appears. Output ONLY the text.",
            image
        ])
        
        if response.text:
            logger.info(f"Gemini Vision extracted {len(response.text)} characters")
            return response.text.strip()
        else:
            return ""
            
    except Exception as e:
        logger.error(f"Gemini Vision OCR failed: {e}")
        return ""

def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from plain text file"""
    try:
        try:
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            text = file_bytes.decode('latin-1')
        return text.strip()
    except Exception as e:
        logger.error(f"Failed to extract text from file: {e}")
        return ""

def extract_text_from_file(file_bytes: bytes, filename: str) -> dict:
    """Extract text from uploaded file based on file type"""
    file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
    logger.info(f"Extracting text from: {filename} ({file_ext})")
    
    extracted_text = ""
    file_type = "UNSUPPORTED"
    
    if file_ext == 'pdf':
        file_type = "PDF"
        extracted_text = extract_text_from_pdf(file_bytes)
    elif file_ext == 'docx':
        file_type = "DOCX"
        extracted_text = extract_text_from_docx(file_bytes)
    elif file_ext in ['txt', 'text', 'log']:
        file_type = "TEXT"
        extracted_text = extract_text_from_txt(file_bytes)
    elif file_ext in ['jpg', 'jpeg', 'png', 'bmp', 'tiff']:
        file_type = "IMAGE"
        extracted_text = extract_text_from_image(file_bytes)
        
    return {
        "filename": filename,
        "file_type": file_type,
        "extracted_text": extracted_text,
        "character_count": len(extracted_text),
        "success": len(extracted_text) > 0
    }
